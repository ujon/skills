#!/usr/bin/env python3
"""
연구노트 DOCX 빌더 — python-docx를 사용한 직접 DOCX 생성

python-docx가 없으면 자동 설치한다.

Usage:
    from docx_builder import DocxNoteBuilder

    doc = DocxNoteBuilder(font_name="맑은 고딕", font_size=11)
    doc.add_header_block("연 구 노 트", "AI 기반 연구")
    doc.add_metadata_table([["연구 제목", "성능 평가"], ["기록자", "홍길동"]])
    doc.add_section(1, "연구 목적 / 배경", "본 연구는...")
    doc.add_signature_block("홍길동", "박검토", "2026-03-13")
    doc.save("./연구노트.docx")
"""

import os
import subprocess
import sys

# ── python-docx 자동 설치 ─────────────────────────────────────

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", "-q", "python-docx"],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

# ── 색상 상수 ──────────────────────────────────────────────────

_HEADER_BG = "1B3A5C"
_TABLE_HEADER_BG = "2E5A88"
_TABLE_BORDER = "B0C4DE"
_SIGNATURE_BG = "F5F8FC"
_META_LABEL_BG = "F0F4F8"
_FOOTER_RGB = RGBColor(0x66, 0x66, 0x66)
_WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
_SUBHEADER_RGB = RGBColor(0xB0, 0xC4, 0xDE)


# ── 유틸 ────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_color: str):
    """셀 배경색 설정."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)


def _set_east_asian_font(run, font_name: str):
    """East Asian 폰트(한글) 설정."""
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>'
        )
        rpr.append(rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font_name)


def _add_run(paragraph, text, font_name, font_size, bold=False, color=None):
    """Run 추가 (폰트·크기·굵기·색상 한 번에 설정)."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    _set_east_asian_font(run, font_name)
    return run


def _set_table_borders(table, hex_color: str):
    """테이블 외곽 및 내부 테두리를 일괄 설정."""
    border_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
        f'</w:tblBorders>'
    )
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(parse_xml(border_xml))


# ── DocxNoteBuilder ─────────────────────────────────────────────

class DocxNoteBuilder:
    """DOCX 빌더."""

    def __init__(self, font_name: str = "맑은 고딕", font_size: int = 11):
        self.doc = Document()
        self.font_name = font_name
        self.font_size = font_size
        self._setup_defaults()

    # ── 초기 설정 ──

    def _setup_defaults(self):
        """기본 스타일·여백 설정."""
        style = self.doc.styles["Normal"]
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} w:eastAsia="{self.font_name}"/>'
            )
            rpr.append(rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), self.font_name)

        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    # ── 헤더 블록 (본문 상단 색상 영역) ──

    def add_header_block(self, title: str, subtitle: str):
        """문서 상단 타이틀 영역 (짙은 남색 배경)."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, _HEADER_BG)

        cell = table.cell(0, 0)
        _set_cell_shading(cell, _HEADER_BG)

        # 제목
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, title, self.font_name, 16, bold=True, color=_WHITE_RGB)

        # 부제(과제명)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p2, subtitle, self.font_name, 11, color=_SUBHEADER_RGB)

        self.doc.add_paragraph()  # 간격

    # ── 메타데이터 테이블 ──

    def add_metadata_table(self, rows: list[list[str]]):
        """2열 메타데이터 테이블 (라벨 | 값)."""
        table = self.doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, _TABLE_BORDER)

        for i, (label, value) in enumerate(rows):
            # 라벨 셀
            c_label = table.cell(i, 0)
            c_label.text = ""
            _set_cell_shading(c_label, _META_LABEL_BG)
            p = c_label.paragraphs[0]
            _add_run(p, label, self.font_name, self.font_size, bold=True)

            # 값 셀
            c_value = table.cell(i, 1)
            c_value.text = ""
            p2 = c_value.paragraphs[0]
            _add_run(p2, value, self.font_name, self.font_size)

        # 열 너비 설정
        for row in table.rows:
            row.cells[0].width = Cm(4)
            row.cells[1].width = Cm(12)

        self.doc.add_paragraph()  # 간격

    # ── 본문 섹션 ──

    def add_section(self, number: int, title: str, content: str):
        """번호 매긴 섹션 (제목 + 본문)."""
        p = self.doc.add_paragraph()
        _add_run(p, f"{number}. {title}", self.font_name, 12, bold=True)

        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p2 = self.doc.add_paragraph()
                _add_run(p2, para_text, self.font_name, self.font_size)

        self.doc.add_paragraph()  # 섹션 간 간격

    # ── 데이터 첨부 ──

    def add_data_table(self, data: list[list], summary: str = ""):
        """데이터 테이블 첨부."""
        if summary:
            p = self.doc.add_paragraph()
            _add_run(p, f"[{summary}]", self.font_name, self.font_size)

        if not data or not data[0]:
            return

        col_cnt = len(data[0])
        table = self.doc.add_table(rows=len(data), cols=col_cnt)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, _TABLE_BORDER)

        for i, row in enumerate(data):
            for j, cell_val in enumerate(row):
                if j >= col_cnt:
                    break
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                is_header = i == 0
                run = _add_run(
                    p, str(cell_val), self.font_name, 9,
                    bold=is_header,
                    color=_WHITE_RGB if is_header else None,
                )
                if is_header:
                    _set_cell_shading(cell, _TABLE_HEADER_BG)

        self.doc.add_paragraph()

    def add_code_block(self, code: str, summary: str = ""):
        """코드 블록 첨부."""
        if summary:
            p = self.doc.add_paragraph()
            _add_run(p, f"[{summary}]", self.font_name, self.font_size)

        for line in code.split("\n")[:50]:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # 배경색
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
            )
            p._element.get_or_add_pPr().append(shading)

    def add_image_ref(self, info: dict):
        """이미지 참조 텍스트 추가."""
        p = self.doc.add_paragraph()
        _add_run(p, f"[이미지: {info.get('path', '')}]", self.font_name, self.font_size)
        dim = ""
        if info.get("width") and info.get("height"):
            dim = f" ({info['width']}×{info['height']})"
        p2 = self.doc.add_paragraph()
        _add_run(
            p2,
            f"  파일: {info.get('ext', '?')}{dim}, {info.get('size_human', '')}",
            self.font_name, self.font_size - 1,
        )

    # ── 서명란 ──

    def add_signature_block(self, 기록자: str, 점검자: str = "", 날짜: str = "____-__-__"):
        """서명란 테이블."""
        self.doc.add_paragraph()

        rows_data = [
            ["구분", "서명", "일자"],
            [f"기록자 ({기록자})", "_________", 날짜],
        ]
        if 점검자:
            rows_data.append([f"점검자 ({점검자})", "_________", "____-__-__"])
        else:
            rows_data.append(["점검자", "_________", "____-__-__"])

        table = self.doc.add_table(rows=len(rows_data), cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, _TABLE_BORDER)

        for i, row in enumerate(rows_data):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                is_header = i == 0
                _add_run(
                    p, val, self.font_name, self.font_size,
                    bold=is_header,
                    color=_WHITE_RGB if is_header else None,
                )
                if is_header:
                    _set_cell_shading(cell, _TABLE_HEADER_BG)
                else:
                    _set_cell_shading(cell, _SIGNATURE_BG)

    # ── 머리글 / 바닥글 ──

    def set_doc_header(self, 과제명: str = "", 기관명: str = ""):
        """페이지 머리글."""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        text = "연 구 노 트"
        if 과제명:
            text += f"  |  {과제명}"

        _add_run(p, text, self.font_name, 8, color=_FOOTER_RGB)

    def set_doc_footer(self, 과제명: str = "", 기관명: str = ""):
        """페이지 바닥글 (기관명 | 과제명 | 페이지번호)."""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        if 기관명:
            parts.append(기관명)
        if 과제명:
            parts.append(과제명)

        if parts:
            _add_run(p, " | ".join(parts) + "  |  ", self.font_name, 8, color=_FOOTER_RGB)

        # 페이지 번호 필드
        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
            f'<w:r><w:rPr><w:sz w:val="16"/></w:rPr><w:t>1</w:t></w:r>'
            f"</w:fldSimple>"
        )
        p._element.append(parse_xml(fld_xml))

    # ── 저장 ──

    def save(self, path: str) -> str:
        """DOCX 파일 저장. 절대 경로를 반환."""
        os.makedirs(os.path.dirname(os.path.abspath(path)) or ".", exist_ok=True)
        self.doc.save(path)
        return os.path.abspath(path)
